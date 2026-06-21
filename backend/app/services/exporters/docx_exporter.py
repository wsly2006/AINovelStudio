"""docx 导出器。

KDP 接受 docx 上传(也接受 epub),docx 在排版控制上更直观。
设计要点:
- 标题样式:章标题 Heading 1
- 章节间分页符,符合 KDP 排版规范
- 段落首行缩进 2 字符(中文习惯)
- blurb / 元数据写到 core_properties,KDP 上传后会读
"""

from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from sqlalchemy.orm import Session

from app.models.project import Project


_PARA_SPLIT = re.compile(r"\n\s*\n+")


def _paragraphs(content: str) -> list[str]:
    out: list[str] = []
    for raw in _PARA_SPLIT.split(content or ""):
        para = raw.strip()
        if para:
            out.append(para)
    return out


def export_to_docx(db: Session, project_id: int) -> bytes:
    p = db.get(Project, project_id)
    if p is None:
        raise ValueError(f"工程不存在: {project_id}")

    doc = Document()

    # core 属性
    cp = doc.core_properties
    cp.title = p.name
    if p.pen_name:
        cp.author = p.pen_name
    if p.blurb:
        cp.comments = p.blurb
    elif p.description:
        cp.comments = p.description
    if p.keywords:
        cp.keywords = ", ".join(p.keywords)
    if p.categories:
        cp.subject = ", ".join(p.categories)

    # 标题页
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(p.name)
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_para.alignment = 1  # CENTER

    if p.pen_name:
        author_para = doc.add_paragraph()
        author_run = author_para.add_run(p.pen_name)
        author_run.font.size = Pt(14)
        author_para.alignment = 1

    if p.blurb:
        doc.add_paragraph()
        for line in p.blurb.splitlines():
            doc.add_paragraph(line)

    # 章节
    chapters = sorted(p.chapters, key=lambda c: c.order_index)
    for idx, ch in enumerate(chapters):
        # 第一章前不加分页符,从第二章起每章前分页(KDP 推荐)
        if idx > 0:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        else:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        sub = (ch.title or "").strip()
        heading_text = f"第 {ch.order_index} 章 {sub}".rstrip()
        doc.add_heading(heading_text, level=1)

        for para_text in _paragraphs(ch.content or ""):
            para = doc.add_paragraph(para_text)
            # 段首缩进
            para.paragraph_format.first_line_indent = Pt(24)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
